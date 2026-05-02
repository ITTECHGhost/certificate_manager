import sys
import docx
from docx import Document

def inspect_docx(filepath):
    doc = Document(filepath)
    print("--- Paragraphs ---")
    for p in doc.paragraphs:
        print(p.text)
            
    print("\n--- Tables ---")
    for i, t in enumerate(doc.tables):
        try:
            for r in t.rows:
                try:
                    row_text = []
                    for c in r.cells:
                        text = c.text.strip().replace('\n', ' ')
                        row_text.append(text)
                    if row_text:
                        print(f"Table {i} Row:", " | ".join(row_text))
                except Exception as e:
                    pass
        except Exception as e:
            print(f"Error reading table {i}: {e}")

if __name__ == '__main__':
    inspect_docx(sys.argv[1])
