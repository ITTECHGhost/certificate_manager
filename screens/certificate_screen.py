import os
import threading
import difflib
import customtkinter as ctk
import win32api
from docxtpl import DocxTemplate
import re
from docx import Document
import jinja2
from config import AppFonts, AppColors
from data.queries import fuzzy_search_students, get_full_certificate_data
from db import get_grade
from ui.base_screen import BaseScreen
from ui.widgets import make_section_header, make_primary_button

class CertificateScreen(BaseScreen):
    """
    Screen for generating and printing student certificates.
    """
    def __init__(self, parent, switch_callback):
        self._selected_student_id = None
        self._templates_dir = "templets"
        super().__init__(parent, switch_callback)

    def _get_available_templates(self) -> list[str]:
        """Scans the templets folder for docx files."""
        if not os.path.exists(self._templates_dir):
            os.makedirs(self._templates_dir, exist_ok=True)
        files = [f for f in os.listdir(self._templates_dir) if f.endswith('.docx') and not f.startswith('~')]
        return files if files else ["لا توجد قوالب (No templates found)"]

    def _build(self) -> None:
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(2, weight=1)

        # --- Top Search Area ---
        top_frame = ctk.CTkFrame(self, fg_color="transparent")
        top_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        top_frame.grid_columnconfigure(0, weight=1)

        make_section_header(top_frame, "إصدار الوثيقة", "Issue Certificate").grid(row=0, column=1, sticky="e")

        search_frame = ctk.CTkFrame(top_frame, fg_color="transparent")
        search_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(10, 4))
        search_frame.grid_columnconfigure(0, weight=1)

        self._search_var = ctk.StringVar()
        self._search_var.trace_add("write", self._on_search_change)

        self._search_entry = ctk.CTkEntry(
            search_frame,
            textvariable=self._search_var,
            placeholder_text="ابحث باسم الطالب (عربي أو إنكليزي)  —  Search student name...",
            font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_BODY),
            height=40, justify="right",
        )
        self._search_entry.grid(row=0, column=0, sticky="ew")

        ctk.CTkButton(
            search_frame, text="بحث\nSearch",
            font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_TINY),
            width=70, height=40, corner_radius=8,
            command=self._do_search,
        ).grid(row=0, column=1, padx=(6, 0))

        # --- Suggestions Frame ---
        self._suggestion_frame = ctk.CTkScrollableFrame(self, height=180, fg_color=("gray94", "gray18"))
        self._suggestion_frame.grid_columnconfigure(0, weight=1)

        # --- Content Area (Left: Student Info, Right: Options) ---
        self._content_frame = ctk.CTkFrame(self, fg_color="transparent")
        self._content_frame.grid(row=2, column=0, columnspan=2, sticky="nsew")
        self._content_frame.grid_columnconfigure(0, weight=1)
        self._content_frame.grid_columnconfigure(1, weight=1)
        self._content_frame.grid_rowconfigure(0, weight=1) # FIX: Makes the boxes fill the screen vertically

        # Left: Student Card
        self._student_card = ctk.CTkFrame(self._content_frame, corner_radius=10, border_width=1, border_color=AppColors.BORDER)
        self._student_card.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        self._student_card.grid_columnconfigure(0, weight=1)
        self._student_card.grid_rowconfigure(1, weight=1)
        
        self._student_info_lbl = ctk.CTkLabel(
            self._student_card,
            text="الرجاء البحث واختيار طالب لعرض معلوماته.\n\nPlease search and select a student.",
            font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_BODY),
            text_color=AppColors.TEXT_MUTED, justify="center"
        )
        self._student_info_lbl.grid(row=0, column=0, pady=20, padx=10)

        self._preview_textbox = ctk.CTkTextbox(
            self._student_card, 
            font=ctk.CTkFont(family="Courier", size=AppFonts.SIZE_BODY),
            state="disabled", wrap="word"
        )
        self._preview_textbox.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        
        # Right: Options Panel
        options_panel = ctk.CTkFrame(self._content_frame, corner_radius=10, border_width=1, border_color=AppColors.BORDER)
        options_panel.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        options_panel.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(
            options_panel, text="خيارات الوثيقة  —  Certificate Options",
            font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_SUBHEADING, weight="bold")
        ).grid(row=0, column=0, pady=(15, 5))

        # Template Selection
        self._template_var = ctk.StringVar()
        self._template_dropdown = ctk.CTkOptionMenu(
            options_panel, variable=self._template_var,
            values=self._get_available_templates(),
            font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_BODY)
        )
        self._template_dropdown.grid(row=1, column=0, sticky="ew", padx=20, pady=(0, 15))

        # Title Input
        title_frame = ctk.CTkFrame(options_panel, fg_color="transparent")
        title_frame.grid(row=2, column=0, sticky="ew", padx=15, pady=5)
        title_frame.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(title_frame, text="إلى / To: ", font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_BODY)).grid(row=0, column=0, sticky="w")
        self._title_entry = ctk.CTkEntry(title_frame, placeholder_text="Whom it May Concern", justify="left")
        self._title_entry.grid(row=0, column=1, sticky="ew", padx=10)

        # Checkboxes and Inputs
        self._opt_summer = ctk.CTkCheckBox(options_panel, text="التدريب الصيفي / Summer Training")
        self._opt_summer.grid(row=3, column=0, sticky="w", padx=20, pady=5)
        self._summer_entry = ctk.CTkEntry(options_panel, placeholder_text="سنة التدريب / Year")
        self._summer_entry.grid(row=4, column=0, sticky="ew", padx=40, pady=(0, 10))

        # Editable Order Number & Date
        self._opt_order = ctk.CTkCheckBox(options_panel, text="الأمر الجامعي / University Order")
        self._opt_order.grid(row=5, column=0, sticky="w", padx=20, pady=5)
        order_frame = ctk.CTkFrame(options_panel, fg_color="transparent")
        order_frame.grid(row=6, column=0, sticky="ew", padx=40, pady=(0, 10))
        order_frame.grid_columnconfigure(0, weight=1)
        order_frame.grid_columnconfigure(1, weight=1)
        self._order_num_entry = ctk.CTkEntry(order_frame, placeholder_text="الرقم / Number")
        self._order_num_entry.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        self._order_date_entry = ctk.CTkEntry(order_frame, placeholder_text="التاريخ / Date (YYYY-MM-DD)")
        self._order_date_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0))

        self._opt_postpone = ctk.CTkCheckBox(options_panel, text="سنوات التأجيل / Postponement")
        self._opt_postpone.grid(row=7, column=0, sticky="w", padx=20, pady=5)
        self._postpone_entry = ctk.CTkEntry(options_panel, placeholder_text="السنوات / Years (e.g. Nill)")
        self._postpone_entry.grid(row=8, column=0, sticky="ew", padx=40, pady=(0, 10))

        self._opt_second_trial = ctk.CTkCheckBox(options_panel, text="الدور الثاني / Second Trial")
        self._opt_second_trial.grid(row=9, column=0, sticky="w", padx=20, pady=5)
        self._second_trial_entry = ctk.CTkEntry(options_panel, placeholder_text="المواد / Subjects (e.g. Nill)")
        self._second_trial_entry.grid(row=10, column=0, sticky="ew", padx=40, pady=(0, 10))

        # Bottom Buttons
        btn_frame = ctk.CTkFrame(options_panel, fg_color="transparent")
        btn_frame.grid(row=11, column=0, sticky="ew", pady=(20, 15))
        btn_frame.grid_columnconfigure(0, weight=1)
        btn_frame.grid_columnconfigure(1, weight=1)

        self._btn_preview = make_primary_button(btn_frame, "📄 معاينة في Word", "Preview", command=self._preview)
        self._btn_preview.grid(row=0, column=0, padx=10, sticky="ew")

        self._btn_print = make_primary_button(btn_frame, "🖨️ طباعة", "Print", command=self._print)
        self._btn_print.grid(row=0, column=1, padx=10, sticky="ew")
        
        self._disable_buttons()

    def refresh(self) -> None:
        # Refresh template list dynamically
        self._template_dropdown.configure(values=self._get_available_templates())
        self._search_var.set("")
        self._hide_suggestions()
        self._selected_student_id = None
        self._student_info_lbl.configure(text="الرجاء البحث واختيار طالب لعرض معلوماته.\n\nPlease search and select a student.")
        self._preview_textbox.configure(state="normal")
        self._preview_textbox.delete("1.0", "end")
        self._preview_textbox.configure(state="disabled")
        self._disable_buttons()

    def _disable_buttons(self):
        self._btn_preview.configure(state="disabled")
        self._btn_print.configure(state="disabled")

    def _enable_buttons(self):
        self._btn_preview.configure(state="normal")
        self._btn_print.configure(state="normal")

    # --- Search Logic ---
    def _on_search_change(self, *_) -> None:
        query = self._search_var.get().strip()
        if len(query) < 2:
            self._hide_suggestions()
            return
        self._show_suggestions_for(query)

    def _do_search(self) -> None:
        query = self._search_var.get().strip()
        if query:
            self._show_suggestions_for(query)

    def _show_suggestions_for(self, query: str) -> None:
        candidates = fuzzy_search_students(query, limit=15)
        if not candidates:
            self._hide_suggestions()
            return

        def similarity(row: dict) -> float:
            ar_score = difflib.SequenceMatcher(None, query, row["full_name_ar"]).ratio()
            en_score = difflib.SequenceMatcher(None, query.lower(), row["full_name_en"].lower()).ratio()
            return max(ar_score, en_score)

        ranked = sorted(candidates, key=similarity, reverse=True)[:8]
        self._render_suggestions(ranked)

    def _render_suggestions(self, students: list[dict]) -> None:
        for w in self._suggestion_frame.winfo_children():
            w.destroy()

        self._suggestion_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 8))
        self._suggestion_frame.tkraise()

        for i, s in enumerate(students):
            dept  = s.get("dept_name_ar", "")
            year  = str(s.get("admission_year", ""))
            label = f"  {s['full_name_ar']}  —  {dept}  |  دفعة {year}"

            ctk.CTkButton(
                self._suggestion_frame,
                text=label,
                font=ctk.CTkFont(family=AppFonts.FAMILY, size=AppFonts.SIZE_SMALL),
                height=36, anchor="e", fg_color="transparent",
                hover_color=AppColors.NAV_HOVER_BG, text_color=AppColors.NAV_TEXT,
                corner_radius=6,
                command=lambda sid=s["id"]: self._select_student(sid),
            ).grid(row=i, column=0, sticky="ew", padx=4, pady=2)

    def _hide_suggestions(self) -> None:
        self._suggestion_frame.grid_remove()
        for w in self._suggestion_frame.winfo_children():
            w.destroy()

    def _select_student(self, student_id: int) -> None:
        self._hide_suggestions()
        self._search_var.set("")
        self._selected_student_id = student_id
        
        try:
            data = get_full_certificate_data(student_id)
            if not data:
                self.show_error("تعذر تحميل بيانات الطالب.")
                return
            
            avg = data.get("average")
            grade_ar, grade_en = get_grade(avg) if avg else ("—", "—")
            rank_text = f"{data['rank']} of {data['total_graduates']}" if data.get("total_graduates") else "—"

            info_text = (
                f"الاسم / Name: {data.get('full_name_en', '')}\n"
                f"القسم / Dept: {data.get('dept_name_en', '')}\n"
                f"المعدل / Average: {avg or '—'}\n"
                f"التقدير / Grade: {grade_en}\n"
                f"الترتيب / Rank: {rank_text}"
            )
            self._student_info_lbl.configure(text=info_text)

            preview_lines = ["--- LIVE PREVIEW / معاينة ---", ""]
            for period in data.get("periods", []):
                preview_lines.append(f"[{period['academic_year']} - Stage {period['stage_number']}]")
                for enr in period.get("enrollments", []):
                    preview_lines.append(f"  • {enr.get('name_en', '')[:30]:<30} | Mark: {enr.get('score', '')} | Unit: {enr.get('credit_hours', '')}")
                preview_lines.append("")
            

# --- REPLACE THIS SECTION IN _select_student ---
            
            self._preview_textbox.configure(state="normal")
            self._preview_textbox.delete("1.0", "end")
            self._preview_textbox.insert("1.0", "\n".join(preview_lines))
            self._preview_textbox.configure(state="disabled")

            # 1. Clear previous text from the input boxes
            self._order_num_entry.delete(0, 'end')
            self._order_date_entry.delete(0, 'end')

            # 2. Retrieve the order data from the database
            order_num = data.get("order_number")
            order_date = data.get("order_date")

            # 3. If the student has an order, check the box and fill the inputs
            if order_num:
                self._opt_order.select()
                self._order_num_entry.insert(0, str(order_num))
                
                if order_date:
                    self._order_date_entry.insert(0, str(order_date))
            else:
                self._opt_order.deselect()

            self._enable_buttons()
            
            # -----------------------------------------------

            # self._preview_textbox.configure(state="normal")
            # self._preview_textbox.delete("1.0", "end")
            # self._preview_textbox.insert("1.0", "\n".join(preview_lines))
            # self._preview_textbox.configure(state="disabled")

            # Reset order entries
            # self._order_num_entry.delete(0, 'end')
            # self._order_date_entry.delete(0, 'end')

            # if data.get("order_number"):
            #     self._opt_order.select()
            #     self._order_num_entry.insert(0, data.get("order_number"))
            #     self._order_date_entry.insert(0, data.get("order_date", ""))
            # else:
            #     self._opt_order.deselect()

            # self._enable_buttons()
        except Exception as e:
            self.show_error(f"Error loading student: {e}")

    # --- Generation Logic ---
    def _build_context(self) -> dict:
        if not self._selected_student_id:
            return {}
        data = get_full_certificate_data(self._selected_student_id)
        if not data:
            return {}

        avg = data.get("average", "")
        grade_ar, grade_en = get_grade(avg) if avg else ("", "")
        
        # 1. Automatic Language Logic based on Template Name
        selected_template = self._template_var.get().lower()
        is_english = "en" in selected_template

        study_type = data.get("study_type", "")
        if is_english:
            study_type_display = "Morning" if study_type == "morning" else ("Evening" if study_type == "evening" else "")
            dept_display = data.get("dept_name_en", "")
            grad_sem_display = "First" if data.get("graduation_semester") == "first" else ("Second" if data.get("graduation_semester") == "second" else "")
            grade_display = grade_en
        else:
            study_type_display = "صباحي" if study_type == "morning" else ("مسائي" if study_type == "evening" else "")
            dept_display = data.get("dept_name_ar", "")
            grad_sem_display = "الاول" if data.get("graduation_semester") == "first" else ("الثاني" if data.get("graduation_semester") == "second" else "")
            grade_display = grade_ar

        # Format semesters
        semesters = []
        for period in data.get("periods", []):
            sem_label = f"Stage {period['stage_number']} - {period['academic_year']}" if is_english else f"المرحلة {period['stage_number']} - {period['academic_year']}"
            doc_rows = []
            enrollments = period.get("enrollments", [])
            
            for i in range(0, len(enrollments), 2):
                left = enrollments[i]
                right = enrollments[i+1] if i+1 < len(enrollments) else {}
                doc_rows.append({
                    "left_name": left.get("name_en" if is_english else "name_ar", ""),
                    "left_mark": str(left.get("score", "")),
                    "left_unit": str(left.get("credit_hours", "")),
                    "right_name": right.get("name_en" if is_english else "name_ar", ""),
                    "right_mark": str(right.get("score", "")),
                    "right_unit": str(right.get("credit_hours", ""))
                })
            
            semesters.append({
                "label": sem_label,
                "rows": doc_rows
            })

        ctx = {
            "Title": self._title_entry.get().strip() or ("Whom it May Concern" if is_english else "إلى من يهمه الأمر"),
            "student_name": data.get("full_name_en" if is_english else "full_name_ar", ""),
            "Birthday": data.get("date_of_birth", ""),    
            "Birthplace": data.get("birthplace_en" if is_english else "birthplace_ar", "") or data.get("birthplace_other", ""),
            "Nationality": data.get("nationality_en" if is_english else "nationality_ar", ""),
            "admission_year": data.get("admission_year", ""),
            "department_id": dept_display,
            "study_type": study_type_display,
            "graduation_date": data.get("graduation_date", ""),
            "graduation_semester": grad_sem_display,
            "average": avg,
            "Grade": grade_display,
            
            # Word's Conditional "If" Triggers
            "sequence_ON": bool(data.get("rank")),
            "Failure_ON": bool(self._opt_postpone.get()),
            "Passed_ON": bool(self._opt_second_trial.get()),

            "Sequence_of_Graduation": data.get("rank", ""),
            "num_students": data.get("total_graduates", ""),
            "Average_of_First_Student": data.get("top_average", ""),
            "semesters": semesters
        }

        # Signatories
        front_sigs = data.get("front_signatories", [])
        for i, sig in enumerate(front_sigs):
            idx = i + 1
            ctx[f"sig{idx}_name"] = sig.get("name_en" if is_english else "name_ar", "")
            ctx[f"sig{idx}_title"] = sig.get("academic_title_en" if is_english else "academic_title_ar", "")
            ctx[f"sig{idx}_resp"] = sig.get("responsibility_en" if is_english else "responsibility_ar", "")
            
        back_sigs = data.get("back_signatories", [])
        for i, sig in enumerate(back_sigs):
            idx = i + 5 
            ctx[f"sig{idx}_name"] = sig.get("name_en" if is_english else "name_ar", "")
            ctx[f"sig{idx}_title"] = sig.get("academic_title_en" if is_english else "academic_title_ar", "")
            ctx[f"sig{idx}_resp"] = sig.get("responsibility_en" if is_english else "responsibility_ar", "")

        # Optionals
        ctx["Summer_Training_year"] = self._summer_entry.get().strip() if self._opt_summer.get() else ""
        
        if self._opt_order.get():
            ctx["order_number"] = self._order_num_entry.get().strip()
            ctx["order_date"] = self._order_date_entry.get().strip()
        else:
            ctx["order_number"] = ""
            ctx["order_date"] = ""

        ctx["Postponement_and_Failure_Years"] = self._postpone_entry.get().strip()
        ctx["Subjects_Passed_with_Second_Trial"] = self._second_trial_entry.get().strip()

        return ctx



    def validate_template_syntax(template_path: str) -> tuple[bool, str]:
        """Scans a Word template and returns the exact visual location of broken tags."""
        try:
            doc = Document(template_path)
        except Exception as e:
            return False, f"Could not read Word file: {e}"

        pattern_start = re.compile(r'\{%[^\s]') 
        pattern_end = re.compile(r'[^\s]%\}')

        # Scan paragraphs
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                if pattern_start.search(p.text) or pattern_end.search(p.text):
                    return False, f"Spacing error in Paragraph {i+1}:\n{p.text}"

        # Scan tables (This tells you the EXACT row)
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        if pattern_start.search(cell.text) or pattern_end.search(cell.text):
                            return False, f"Spacing error in Table {t_idx+1}, Row {r_idx+1}, Column {c_idx+1}:\n{cell.text}"

        return True, "No syntax errors found."

    def _generate_docx(self, output_path: str) -> bool:
        try:
            selected = self._template_var.get()
            if not selected or "لا توجد" in selected:
                self.show_error("الرجاء تحديد قالب صالح / Please select a valid template")
                return False

            template_path = os.path.join(self._templates_dir, selected)
            if not os.path.exists(template_path):
                self.show_error(f"Template not found at: {template_path}")
                return False

            # Require order number if checked
            if self._opt_order.get() and not self._order_num_entry.get().strip():
                self.show_error("الرجاء إدخال رقم الأمر الجامعي / Please enter the Order Number")
                return False

            # ---> NEW: Run the visual pre-flight check <---
            is_valid, error_msg = validate_template_syntax(template_path)
            if not is_valid:
                self.show_error(f"Template Syntax Error:\n{error_msg}")
                return False

            doc = DocxTemplate(template_path)
            ctx = self._build_context()
            
            # ---> NEW: Catch specific Jinja parsing errors <---
            try:
                doc.render(ctx)
            except jinja2.exceptions.TemplateSyntaxError as jinja_err:
                self.show_error(f"Jinja Code Error:\n{jinja_err.message}\n(Likely hidden Word formatting)")
                return False

            doc.save(output_path)
            return True

        except Exception as e:
            self.show_error(f"Error generating document:\n{e}")
            return False

    def _preview(self):
        self._disable_buttons()
        threading.Thread(target=self._run_preview, daemon=True).start()

    def _run_preview(self):
        out_path = os.path.abspath("temp_certificate.docx")
        if self._generate_docx(out_path):
            try:
                os.startfile(out_path)
            except Exception as e:
                self.after(0, lambda: self.show_error(f"Could not open file: {e}"))
        self.after(0, self._enable_buttons)

    def _print(self):
        self._disable_buttons()
        threading.Thread(target=self._run_print, daemon=True).start()

    def _run_print(self):
        out_path = os.path.abspath("temp_certificate.docx")
        if self._generate_docx(out_path):
            try:
                win32api.ShellExecute(0, "print", out_path, None, ".", 0)
            except Exception as e:
                self.after(0, lambda: self.show_error(f"Could not print file: {e}"))
        self.after(0, self._enable_buttons)