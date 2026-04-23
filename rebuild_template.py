"""
rebuild_template.py
===================
One-shot script: rewrites  templets/semster - Temp.docx  so that:
  1. Tables 2–8 (hardcoded sample data) are REMOVED.
  2. Table 1 (the template semester table) is rebuilt with proper docxtpl tags.
  3. Table 2 (front-page signatories) cells are made fully dynamic.
  4. Table 3 (back-page signatories) cells are made fully dynamic.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree
import copy, re

SRC  = r"templets/semster - Temp.bak"
DEST = r"templets/semster - Temp.docx"   # overwrite

def make_tr_with_single_cell_text(text, num_cols=6):
    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    tcp = OxmlElement("w:tcPr")
    gs  = OxmlElement("w:gridSpan")
    gs.set(qn("w:val"), str(num_cols))
    tcp.append(gs)
    tc.append(tcp)
    p  = OxmlElement("w:p")
    r  = OxmlElement("w:r")
    t  = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    tr.append(tc)
    for _ in range(num_cols - 1):
        ghost_tc = OxmlElement("w:tc")
        ghost_tcp = OxmlElement("w:tcPr")
        ghost_vmerge = OxmlElement("w:vMerge")
        ghost_tcp.append(ghost_vmerge)
        ghost_tc.append(ghost_tcp)
        ghost_p = OxmlElement("w:p")
        ghost_tc.append(ghost_p)
        tr.append(ghost_tc)
    return tr

def make_data_tr(c0, c1, c2, c3, c4, c5):
    tr = OxmlElement("w:tr")
    for txt in [c0, c1, c2, c3, c4, c5]:
        tc = OxmlElement("w:tc")
        p  = OxmlElement("w:p")
        r  = OxmlElement("w:r")
        t  = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = txt
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
    return tr

doc = Document(SRC)

# Clean up any bad variable tags in the entire document before proceeding
replacements = {
    'Average of First Student': 'Average_of_First_Student',
    'Semester _ year': 'Semester_year',
    'Sequence of Graduation': 'Sequence_of_Graduation',
    'Postponement and Failure Years': 'Postponement_and_Failure_Years',
    'Subjects Passed with Second Trial': 'Subjects_Passed_with_Second_Trial',
    'Summer Training year': 'Summer_Training_year',
    'Birthday }': 'Birthday',
    'Birthplace }': 'Birthplace',
    'num_students ': 'num_students',
    '{%tr for row in t.rows %}{{row. Course}}': '', # Remove whatever the user did manually
    '{{Unit}} {%tr endfor %}': '',
}

def clean_text(text):
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

for para in doc.paragraphs:
    para.text = clean_text(para.text)

# Step 2: Remove tables 2 through 8 if they still exist (7 tables)
# In my previous run they might have already been removed, so we check.
if len(doc.tables) > 5:
    tables_to_remove = doc.tables[2:9]
    for tbl in tables_to_remove:
        tbl._element.getparent().remove(tbl._element)

doc.save(DEST)
doc = Document(DEST)

# Step 3: Rebuild Table 1 (Semester table)
sem_table = doc.tables[1]
tbl_el    = sem_table._element

for row in list(sem_table.rows):
    tbl_el.remove(row._element)

# Ensure no weird spaces inside docxtpl control tags
tbl_el.append(make_tr_with_single_cell_text("{%tr for sem in semesters %}"))
tbl_el.append(make_data_tr("{{ sem.label }}", "{{ sem.label }}", "{{ sem.label }}", "{{ sem.label }}", "{{ sem.label }}", "{{ sem.label }}"))
tbl_el.append(make_data_tr("Subject", "Mark", "Unit", "Subject", "Mark", "Unit"))
tbl_el.append(make_tr_with_single_cell_text("{%tr for row in sem.rows %}"))
tbl_el.append(make_data_tr("{{ row.left_name }}", "{{ row.left_mark }}", "{{ row.left_unit }}", "{{ row.right_name }}", "{{ row.right_mark }}", "{{ row.right_unit }}"))
tbl_el.append(make_tr_with_single_cell_text("{%tr endfor %}"))
tbl_el.append(make_tr_with_single_cell_text("{%tr endfor %}"))

# Step 4: Signatories (Front)
sig_table = doc.tables[2]
def set_sig_cell(cell, title_var, name_var, resp_var):
    tc = cell._tc
    for p in list(tc.findall(f".//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}p")):
        tc.remove(p)
    p1 = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t1.text = f"{title_var} {name_var}"
    r1.append(t1)
    p1.append(r1)
    tc.append(p1)
    p2 = OxmlElement("w:p")
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t2.text = resp_var
    r2.append(t2)
    p2.append(r2)
    tc.append(p2)

set_sig_cell(sig_table.rows[0].cells[0], "{{sig3_title}}", "{{sig3_name}}", "{{sig3_resp}}")
set_sig_cell(sig_table.rows[0].cells[2], "{{sig1_title}}", "{{sig1_name}}", "{{sig1_resp}}")
if len(sig_table.rows) > 1:
    set_sig_cell(sig_table.rows[1].cells[0], "{{sig4_title}}", "{{sig4_name}}", "{{sig4_resp}}")
    set_sig_cell(sig_table.rows[1].cells[2], "{{sig2_title}}", "{{sig2_name}}", "{{sig2_resp}}")

# Step 5: Signatories (Back)
back_table = doc.tables[3]
def set_back_sig_cell(cell, name_var, resp_var):
    tc = cell._tc
    for p in list(tc.findall(f".//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}p")):
        tc.remove(p)
    p1 = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    t1 = OxmlElement("w:t")
    t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t1.text = name_var
    r1.append(t1)
    p1.append(r1)
    tc.append(p1)
    p2 = OxmlElement("w:p")
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:t")
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t2.text = resp_var
    r2.append(t2)
    p2.append(r2)
    tc.append(p2)

set_back_sig_cell(back_table.rows[0].cells[0], "{{sig6_name}}", "{{sig6_resp}}")
set_back_sig_cell(back_table.rows[0].cells[1], "{{sig5_name}}", "{{sig5_resp}}")

doc.save(DEST)
print("Template rebuilt.")
