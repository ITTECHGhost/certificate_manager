import docx

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

def fix_runs(blocks):
    for block in blocks:
        for p in block.paragraphs:
            runs = p.runs
            for i in range(1, len(runs) - 1):
                if runs[i].text == ' ':
                    if runs[i-1].text.endswith('{%') and runs[i+1].text.startswith('tr'):
                        runs[i].text = ''
                        print("Fixed '{% tr' -> '{%tr'")
                    elif runs[i-1].text.endswith('{%') and runs[i+1].text.startswith('p '):
                        runs[i].text = ''
                        print("Fixed '{% p ' -> '{%p '")
                    elif runs[i-1].text.endswith('{%') and runs[i+1].text.startswith('tc '):
                        runs[i].text = ''
                        print("Fixed '{% tc ' -> '{%tc '")
                        
            # Also fix 'endif%}' to 'endif %}'
            for r in runs:
                if r.text == 'endif':
                    # Check next run
                    pass
                if 'endif%}' in r.text:
                    r.text = r.text.replace('endif%}', 'endif %}')
                if 'Passed_ON%}' in r.text:
                    r.text = r.text.replace('Passed_ON%}', 'Passed_ON %}')

for t in doc.tables:
    for row in t.rows:
        fix_runs(row.cells)

fix_runs([doc])

doc.save(r'templets\semester - Temp - En - D_fixed2.docx')
print("Run fixer complete.")
