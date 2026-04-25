import docx
import re

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

for t in doc.tables:
    for r in t.rows:
        text = ''.join(c.text for c in r.cells)
        if 'sequence_ON' in text:
            # We found the row!
            # Let's remove {%tr if sequence_ON %} and {%tr endif %} from ALL cells
            for c in r.cells:
                for p in c.paragraphs:
                    if '{%tr if sequence_ON %}' in p.text:
                        p.text = p.text.replace('{%tr if sequence_ON %}', '')
                    if '{%tr endif %}' in p.text:
                        p.text = p.text.replace('{%tr endif %}', '')
            
            # Now, prepend the tags to the first cell's first paragraph
            if r.cells:
                first_p = r.cells[0].paragraphs[0]
                first_p.text = '{%tr if sequence_ON %}' + first_p.text + '{%tr endif %}'
            
            break

doc.save(r'templets\semester - Temp - En - D.docx')
print("Fixed sequence_ON tags in DOCX.")
