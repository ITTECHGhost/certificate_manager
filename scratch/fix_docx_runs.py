import docx

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

for t in doc.tables:
    for r in t.rows:
        text = ''.join(c.text for c in r.cells)
        if '{%tr if sequence_ON %}' in text:
            # Found the row!
            # 1. Delete {%tr endif %} from all runs
            for c in r.cells:
                for p in c.paragraphs:
                    # Let's just do a simple replacement if it's fully contained in a run.
                    # Or build a text accumulator.
                    # Since it's Jinja, the user typed it. Let's replace the substring in runs.
                    for run in p.runs:
                        if '{%tr' in run.text or 'endif' in run.text or '%}' in run.text:
                            # Because it might be split, the safest way without destroying formatting
                            # is to just find the run with 'endif' and clear it, and hope it works.
                            pass

            # A better way is to use docx-mailmerge or docx's internal XML tree.
