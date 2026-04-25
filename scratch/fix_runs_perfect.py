import docx

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

def fix_runs(blocks):
    for block in blocks:
        for p in block.paragraphs:
            runs = p.runs
            for i in range(len(runs)):
                if runs[i].text == ' ':
                    prev_text = ''.join(r.text for r in runs[:i])
                    next_text = ''.join(r.text for r in runs[i+1:])
                    if prev_text.endswith('{%') and (next_text.startswith('tr') or next_text.startswith('p') or next_text.startswith('tc')):
                        runs[i].text = ''
                        print("Fixed space between {% and tr/p/tc")
                
                # Also fix endif%}
                if 'endif%}' in runs[i].text:
                    runs[i].text = runs[i].text.replace('endif%}', 'endif %}')
                if 'Passed_ON%}' in runs[i].text:
                    runs[i].text = runs[i].text.replace('Passed_ON%}', 'Passed_ON %}')
                    
            for r in runs:
                if '{% tr' in r.text:
                    r.text = r.text.replace('{% tr', '{%tr')
                if '{% p' in r.text:
                    r.text = r.text.replace('{% p', '{%p')
                if '{% tc' in r.text:
                    r.text = r.text.replace('{% tc', '{%tc')
                if '% tr' in r.text:
                    r.text = r.text.replace('% tr', '%tr')
                if '% p' in r.text:
                    r.text = r.text.replace('% p', '%p')
                if '% tc' in r.text:
                    r.text = r.text.replace('% tc', '%tc')
                    
            for i in range(len(runs) - 1):
                if runs[i].text.endswith('{% ') and runs[i+1].text.startswith('tr'):
                    runs[i].text = runs[i].text[:-1]

for t in doc.tables:
    for row in t.rows:
        fix_runs(row.cells)

fix_runs([doc])

doc.save(r'templets\semester - Temp - En - D.docx')
print("Runs fixed safely part 2!")
