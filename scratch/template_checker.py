import re
import sys
try:
    from docx import Document
except ImportError:
    print("Error: The 'python-docx' library is not installed.")
    print("Please run: pip install python-docx")
    sys.exit(1)

def check_docxtpl_template(file_path):
    print(f"--- Starting Diagnostic Check for: {file_path} ---")
    
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Failed to open document: {e}")
        return

    errors_found = 0
    
    # Extract all text blocks from paragraphs and tables
    text_blocks = []
    
    # 1. Get paragraph text
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            text_blocks.append((f"Paragraph {i+1}", p.text))
            
    # 2. Get table cell text
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    text_blocks.append((f"Table {t_idx+1}, Row {r_idx+1}, Col {c_idx+1}", cell.text))

    # Regex patterns for common docxtpl errors
    # Looks for {% missing a space right after it (e.g. {%tr)
    pattern_missing_start_space = re.compile(r'\{%[^\s]') 
    # Looks for %} missing a space right before it (e.g. ON%})
    pattern_missing_end_space = re.compile(r'[^\s]%\}')
    # Captures everything inside {% ... %} to check inner syntax
    pattern_all_tags = re.compile(r'\{%(.*?)%\}')

    for location, text in text_blocks:
        # Check 1: Missing space after {%
        if pattern_missing_start_space.search(text):
            print(f"\n❌ ERROR at [{location}]")
            print(f"  Issue: Missing space after '{{%'")
            print(f"  Text:  {text.strip()}")
            errors_found += 1

        # Check 2: Missing space before %}
        if pattern_missing_end_space.search(text):
            print(f"\n❌ ERROR at [{location}]")
            print(f"  Issue: Missing space before '%}}'")
            print(f"  Text:  {text.strip()}")
            errors_found += 1

        # Check 3: Check inner tag spelling/syntax
        tags = pattern_all_tags.findall(text)
        for tag in tags:
            tag_clean = tag.strip()
            
            # Check for accidental spaces in keywords (e.g. "end if" instead of "endif")
            if "end if" in tag_clean:
                print(f"\n❌ ERROR at [{location}]")
                print(f"  Issue: Typo in tag. 'end if' should be 'endif'")
                print(f"  Text:  {{%{tag}%}}")
                errors_found += 1
                
            if "end for" in tag_clean:
                print(f"\n❌ ERROR at [{location}]")
                print(f"  Issue: Typo in tag. 'end for' should be 'endfor'")
                print(f"  Text:  {{%{tag}%}}")
                errors_found += 1

    print("\n--------------------------------------------------")
    if errors_found == 0:
        print("✅ Check Complete: No obvious spacing or syntax errors found in tags!")
        print("   (If it still crashes, Word may be hiding XML formatting inside the tags.)")
    else:
        print(f"🚨 Check Complete: Found {errors_found} potential error(s). Please fix them in Word.")

if __name__ == "__main__":
    # ---> CHANGE THIS TO YOUR ACTUAL TEMPLATE NAME <---
    TEMPLATE_PATH = "templets/semester - Temp - En - D.docx" 
    
    import os
    if os.path.exists(TEMPLATE_PATH):
        check_docxtpl_template(TEMPLATE_PATH)
    else:
        print(f"File not found: {TEMPLATE_PATH}")
        print("Make sure the path is correct relative to where you are running this script.")