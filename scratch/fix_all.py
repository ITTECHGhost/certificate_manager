import docx

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

def copy_format(src_run, dst_run):
    if src_run:
        dst_run.font.name = src_run.font.name
        dst_run.font.size = src_run.font.size
        dst_run.font.bold = src_run.font.bold
        dst_run.font.italic = src_run.font.italic
        dst_run.font.color.rgb = src_run.font.color.rgb

for t in doc.tables:
    for row in t.rows:
        text = ''.join(c.text for c in row.cells)
        if 'sequence_ON' in text:
            # Rebuild Cell 0
            fmt_run0 = row.cells[0].paragraphs[0].runs[-1] if row.cells[0].paragraphs[0].runs else None
            row.cells[0].text = ''
            p0 = row.cells[0].paragraphs[0]
            new_run0 = p0.add_run('{%tr if sequence_ON %}{%tr endif %} Sequence_of_Graduation: {{Sequence_of_Graduation}}')
            copy_format(fmt_run0, new_run0)
            
            # Rebuild Cell 2
            fmt_run2 = row.cells[2].paragraphs[0].runs[0] if row.cells[2].paragraphs[0].runs else None
            row.cells[2].text = ''
            p2 = row.cells[2].paragraphs[0]
            new_run2 = p2.add_run('Average_of_First_Student: {{Average_of_First_Student}}')
            copy_format(fmt_run2, new_run2)
            
            # Rebuild Cell 3
            if len(row.cells) > 3:
                fmt_run3 = row.cells[3].paragraphs[0].runs[0] if row.cells[3].paragraphs[0].runs else None
                row.cells[3].text = ''
                p3 = row.cells[3].paragraphs[0]
                new_run3 = p3.add_run('Average_of_First_Student: {{Average_of_First_Student}}')
                copy_format(fmt_run3, new_run3)

        # Fix other spaces
        if 'period.rows' in text:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if '{% tr' in r.text:
                            r.text = r.text.replace('{% tr', '{%tr')
                        if '% tr' in r.text:
                            r.text = r.text.replace('% tr', '%tr')
                            
        # Fix endif%}
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    if 'endif%}' in r.text:
                        r.text = r.text.replace('endif%}', 'endif %}')
                    if 'Passed_ON%}' in r.text:
                        r.text = r.text.replace('Passed_ON%}', 'Passed_ON %}')

# Fix body spaces
for p in doc.paragraphs:
    for r in p.runs:
        if '{% tr' in r.text:
            r.text = r.text.replace('{% tr', '{%tr')
        if '% tr' in r.text:
            r.text = r.text.replace('% tr', '%tr')
        if 'endif%}' in r.text:
            r.text = r.text.replace('endif%}', 'endif %}')
        if 'Passed_ON%}' in r.text:
            r.text = r.text.replace('Passed_ON%}', 'Passed_ON %}')

# One more pass to fix isolated spaces in the document
for block in [doc] + [c for t in doc.tables for r in t.rows for c in r.cells]:
    for p in block.paragraphs:
        for i in range(1, len(p.runs) - 1):
            if p.runs[i].text == ' ':
                prev = ''.join(r.text for r in p.runs[:i])
                nxt = ''.join(r.text for r in p.runs[i+1:])
                if prev.endswith('{%') and (nxt.startswith('tr') or nxt.startswith('tc') or nxt.startswith('p')):
                    p.runs[i].text = ''

doc.save(r'templets\semester - Temp - En - D.docx')
print("Complete fix applied via python-docx!")
