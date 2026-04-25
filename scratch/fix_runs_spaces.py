import docx

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

def fix_runs(blocks):
    for block in blocks:
        for p in block.paragraphs:
            for i, r in enumerate(p.runs):
                if '{% ' in r.text:
                    r.text = r.text.replace('{% ', '{%')
                if '{%  ' in r.text:
                    r.text = r.text.replace('{%  ', '{%')
                if '% tr' in r.text:
                    r.text = r.text.replace('% tr', '%tr')
                
                # Check for split '{%' and ' tr'
                if r.text.endswith('{%'):
                    if i + 1 < len(p.runs) and p.runs[i+1].text.startswith(' tr'):
                        p.runs[i+1].text = p.runs[i+1].text.replace(' tr', 'tr', 1)

# Fix in tables
for t in doc.tables:
    for row in t.rows:
        fix_runs(row.cells)

# Fix in document body
fix_runs([doc])

# Also let's fix endif%} -> endif %}
def fix_endif(blocks):
    for block in blocks:
        for p in block.paragraphs:
            for r in p.runs:
                if 'endif%}' in r.text:
                    r.text = r.text.replace('endif%}', 'endif %}')
                if 'Passed_ON%}' in r.text:
                    r.text = r.text.replace('Passed_ON%}', 'Passed_ON %}')

for t in doc.tables:
    for row in t.rows:
        fix_endif(row.cells)
fix_endif([doc])

doc.save(r'templets\semester - Temp - En - D.docx')
print("Fixed spaces in runs!")
