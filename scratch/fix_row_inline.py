import docx

doc_path = r'templets\semester - Temp - En - D.docx'
doc = docx.Document(doc_path)

def fix_runs(blocks):
    for block in blocks:
        for p in block.paragraphs:
            for r in p.runs:
                if '{%tr if' in r.text:
                    r.text = r.text.replace('{%tr if', '{% if')
                if '{%tr endif' in r.text:
                    r.text = r.text.replace('{%tr endif', '{% endif')

for t in doc.tables:
    for row in t.rows:
        fix_runs(row.cells)

fix_runs([doc])

doc.save(r'templets\semester - Temp - En - D.docx')
print("Replaced {%tr if sequence_ON %} with inline {% if sequence_ON %}!")
